#!/usr/bin/env python3
"""Convert a resume markdown file to a professionally formatted .docx file."""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


FONT_NAME = "Calibri"
FONT_COLOR = RGBColor(0x2D, 0x2D, 0x2D)
HEADING_COLOR = RGBColor(0x1A, 0x1A, 0x1A)
SUBTLE_COLOR = RGBColor(0x55, 0x55, 0x55)


def set_run_font(run, size, bold=False, italic=False, color=None, name=FONT_NAME):
    """Apply font formatting to a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color or FONT_COLOR
    # Ensure Calibri works on East Asian text too
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph spacing in points."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def add_horizontal_rule(doc):
    """Add a thin horizontal rule to the document."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=4)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_formatted_text(paragraph, text, base_size=10.5):
    """Parse inline markdown (bold, italic) and add runs to a paragraph."""
    # Split on bold markers (**text**) and italic markers (*text*)
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            set_run_font(run, base_size, bold=True)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            set_run_font(run, base_size, italic=True, color=SUBTLE_COLOR)
        else:
            if part:
                run = paragraph.add_run(part)
                set_run_font(run, base_size)


def parse_markdown(md_text):
    """Parse markdown into a list of structured blocks."""
    lines = md_text.split("\n")
    blocks = []
    i = 0
    # Track whether we're in the contact header (between H1 and first HR)
    in_contact_block = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", stripped):
            blocks.append({"type": "hr"})
            if in_contact_block:
                in_contact_block = False
            i += 1
            continue

        # H1 heading
        if stripped.startswith("# ") and not stripped.startswith("## "):
            blocks.append({"type": "h1", "text": stripped[2:].strip()})
            in_contact_block = True
            i += 1
            continue

        # In contact block: each line is its own contact line (don't merge)
        if in_contact_block:
            # Skip "## Contact" heading if present — it's structural, not content
            if stripped.startswith("## "):
                i += 1
                continue
            # Strip bold markers from contact lines (e.g., **Name**)
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
            # Skip lines that just repeat the H1 name
            if clean == blocks[-1]["text"] if blocks and blocks[-1]["type"] == "h1" else False:
                i += 1
                continue
            blocks.append({"type": "contact", "text": clean})
            i += 1
            continue

        # H3 heading (check before H2)
        if stripped.startswith("### "):
            blocks.append({"type": "h3", "text": stripped[4:].strip()})
            i += 1
            continue

        # H2 heading
        if stripped.startswith("## "):
            blocks.append({"type": "h2", "text": stripped[3:].strip()})
            i += 1
            continue

        # Bullet list item
        if re.match(r"^[-*]\s", stripped):
            bullet_text = re.sub(r"^[-*]\s+", "", stripped)
            blocks.append({"type": "bullet", "text": bullet_text})
            i += 1
            continue

        # Regular paragraph
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith("#") and not re.match(r"^[-*]\s", lines[i].strip()) and not re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", lines[i].strip()):
            para_lines.append(lines[i].strip())
            i += 1
        blocks.append({"type": "paragraph", "text": " ".join(para_lines)})

    return blocks


def build_docx(blocks, output_path):
    """Build a .docx document from parsed blocks."""
    doc = Document()

    # Set default document margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    for i, block in enumerate(blocks):
        btype = block["type"]

        if btype == "h1":
            # Name / Title - large centered
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=0, after=2)
            run = p.add_run(block["text"])
            set_run_font(run, 22, bold=True, color=HEADING_COLOR)
            continue

        if btype == "contact":
            # Contact info lines - centered, subtle, tight spacing
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=0, after=0, line=14)
            run = p.add_run(block["text"])
            set_run_font(run, 9.5, color=SUBTLE_COLOR)
            continue

        if btype == "hr":
            add_horizontal_rule(doc)
            continue

        if btype == "h2":
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=10, after=3)
            run = p.add_run(block["text"].upper())
            set_run_font(run, 11.5, bold=True, color=HEADING_COLOR)
            continue

        if btype == "h3":
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=6, after=1)
            add_formatted_text(p, block["text"], base_size=10.5)
            # Make the whole thing bold if not already formatted
            if "**" not in block["text"]:
                for run in p.runs:
                    run.font.bold = True
            continue

        if btype == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, before=1, after=1, line=14)
            # Clear default text and use formatted version
            p.clear()
            add_formatted_text(p, block["text"], base_size=10)
            # Adjust indent
            pf = p.paragraph_format
            pf.left_indent = Inches(0.3)
            pf.first_line_indent = Inches(-0.15)
            continue

        if btype == "paragraph":
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=2, after=2, line=15)
            add_formatted_text(p, block["text"], base_size=10.5)
            continue

    doc.save(str(output_path))


def main():
    parser = argparse.ArgumentParser(description="Convert resume markdown to docx")
    parser.add_argument("input", help="Path to the markdown resume file")
    parser.add_argument("output", help="Path for the output .docx file")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)

    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    md_text = input_path.read_text(encoding="utf-8")
    blocks = parse_markdown(md_text)
    build_docx(blocks, output_path)
    print(f"Created: {output_path}")


if __name__ == "__main__":
    main()
